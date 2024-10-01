import pandas as pd
from io import BytesIO
import os
import sys
from datetime import datetime
from docx import Document
from docx.oxml import OxmlElement, qn
import shutil
from docx.enum.text import WD_BREAK
from docx.shared import Inches
import os
import shutil
from docxcompose.composer import Composer



sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from src.utils import load_config, load_translations
from src.load_data import get_project_path, load_file, get_download_link


def create_titel_page(version, language_code):
    organisation_folder = get_project_path('organisation_data')
    translations = load_translations(organisation_folder / 'translations.json')

    titel = translations['word_report']['titel'][language_code]

    today_date = datetime.today().strftime('%Y-%m-%d')
    
    # Prepare the data to fill in the template
    data = {
        "Titel": titel,
        "Date": today_date,
        "Version": version,
    }

    template_path = 'organisation_data/template_titelpage.docx'
    output_path = 'organisation_data/populated_titelpage.docx'

    populate_template(template_path, data, output_path)

# --- Working ---

def insert_page_break(doc):
    """Inserts a page break by modifying the XML structure."""
    # Create the <w:br> element with the "page" type
    run_element = doc.paragraphs[-1]._p  # Get the last paragraph's XML element
    br = OxmlElement("w:br")
    br.set(qn('w:type'), "page")
    run_element.append(br)

def merge_documents(output_path, temp_docs):
    """Merges a list of temp documents into the final output document, adding a page break between each."""
    # Create an initial document
    master = Document(temp_docs[0])
    composer = Composer(master)

    # Append the rest of the documents, adding a page break after each
    for temp_doc_path in temp_docs[1:]:
        temp_doc = Document(temp_doc_path)
        insert_page_break(master)  # Insert a page break before appending each document
        composer.append(temp_doc)

    # Save the final merged document
    composer.save(output_path)

def create_element_overview(version, language_code):
    organisation_folder = get_project_path('organisation_data')
    translations = load_translations(organisation_folder / 'translations.json')

    titel = translations['word_report']['titel'][language_code]
    today_date = datetime.today().strftime('%Y-%m-%d')

    # Load the workflow data from the CSV file
    element_df = load_file(version, 'M_Elements.csv')


    # Define the template path
    template_path = 'organisation_data/template_element.docx'

    # Ensure the temp folder exists
    temp_folder = 'organisation_data/temp/'
    if not os.path.exists(temp_folder):
        os.makedirs(temp_folder)

    # List to store paths of temporary files
    temp_docs = []

    element_df = element_df.reset_index(drop=True)
    print(element_df)

    # Loop through each row in the workflow DataFrame
    for index, row in element_df.iterrows():
        # Prepare the data to fill in the template
        print(f"Processing row {index}: {row[f'ElementName{language_code}']}")
        data = {
            "Titel": titel,
            "Date": today_date,
            "Version": version,
            "ElementName": row[f'ElementName{language_code}'],
            "IfcEntity": row[f'IfcEntityIfc4.0Name'],
            "ElementDescription": row[f'ElementDescription{language_code}'],
        }

        
        picture_name = row['ImageName']

        if pd.isna(picture_name):
            image_path = None
        else:
            image_path = get_download_link(version=version, file_name=picture_name, data_folder='data')

        # Create a temporary output path for each row's document
        output_path_temp = f'organisation_data/temp/{index}.docx'

        # Copy the template to the temp file location to ensure formatting is preserved
        try:
            shutil.copy(template_path, output_path_temp)
        except FileNotFoundError as e:
            print(f"Error copying template: {e}")
            continue  # Skip to the next row if there's an error

        # Populate the template with the data and save to temp file
        populate_template(output_path_temp, data, output_path_temp, image_key='ElementPicture', image_path=image_path)

        # Add the path to the list of temp docs
        temp_docs.append(output_path_temp)

    # Merge all temp documents into the final output document
    final_output_path = 'organisation_data/populated_element_overview.docx'
    merge_documents(final_output_path, temp_docs)

    print(f"Populated document saved at {final_output_path}")
    


def populate_template(template_path, data, output_path, image_key="Image", image_path=None):
    """Populate the Word template with data and replace {Image} with an image or handle placeholders in the footer."""
    doc = Document(template_path)

    # Replace placeholders in the main body paragraphs
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            value = str(value)  # Ensure the value is always a string
            
            if f'{{{key}}}' in paragraph.text:
                paragraph.text = paragraph.text.replace(f'{{{key}}}', value)

        # Handle image placeholders
        if image_key in paragraph.text:
            if image_path:
                # Clear the paragraph text first and insert the image
                paragraph.text = ""
                run = paragraph.add_run()
                run.add_picture(image_path, width=Inches(6))  # Adjust size if needed
            else:
                # Remove the placeholder by replacing it with an empty string
                paragraph.text = paragraph.text.replace(f'{{{image_key}}}', "")

    # Now, handle placeholders in the footer
    for section in doc.sections:
        footer = section.footer
        for paragraph in footer.paragraphs:
            for key, value in data.items():
                value = str(value)
                if f'{{{key}}}' in paragraph.text:
                    paragraph.text = paragraph.text.replace(f'{{{key}}}', value)

            # Handle image placeholders in the footer
            if image_key in paragraph.text:
                if image_path:
                    paragraph.text = ""
                    run = paragraph.add_run()
                    run.add_picture(image_path, width=Inches(6))
                else:
                    paragraph.text = paragraph.text.replace(f'{{{image_key}}}', "")

    # Save the updated document
    doc.save(output_path)
    return(output_path)

### --- works ---

def create_usecase_table(version, language_code):
    organisation_folder = get_project_path('organisation_data')
    translations = load_translations(organisation_folder / 'translations.json')

    today_date = datetime.today().strftime('%Y-%m-%d')

    # Load the workflow data from the CSV file
    workflow_df = load_file(version, 'M_Workflows.csv')

    # Define the template path
    template_path = 'organisation_data/template_usecase.docx'
    doc = Document(template_path)

    # Find the table that contains the placeholders {Code}, {WorkflowName}, and {WorkflowDescription}
    table = doc.tables[0]  # Assuming the template only has one table, or index it correctly

    # Loop through the DataFrame and add a new row for each entry
    for index, row in workflow_df.iterrows():
        # Add a new row to the table
        new_row = table.add_row()

        # Populate the new row with data from the DataFrame
        new_row.cells[0].text = str(row['WorkflowCode'])
        new_row.cells[1].text = str(row[f'WorkflowName{language_code}'])
        new_row.cells[2].text = str(row[f'WorkflowDescription{language_code}'])

        # Optional: Set font size for the table cells if needed
        for cell in new_row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)

    # Save the updated document
    output_path = 'organisation_data/populated_usecase_table.docx'
    doc.save(output_path)
    return(output_path)


titel_doc = create_titel_page('V1', 'DE')

usecase_doc = create_usecase_table('V1', 'DE')

element_doc = create_element_overview('V1', 'DE')


